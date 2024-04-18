import os
import re
import docx
import pdfplumber
import fitz
from paddleocr import PaddleOCR
from pathlib import Path
from bs4 import BeautifulSoup
import requests
import pnlp
import pandas as pd
import subprocess

title_reg = re.compile(r"^第[一二三四五六七八九十 ]+[章条]+")


class ExcelReader:

    def __init__(self):
        ...

    def convert(self, file_path: str, save_path: str) -> str:
        df = pd.read_excel(file_path)
        text = "一、"
        text += " ".join(df.columns)
        title = Path(file_path).stem
        out_fn = title + ".txt"
        out_fp = os.path.join(save_path, out_fn)
        with open(out_fp, "w") as f:
            f.write(text)
        return out_fn


class HtmlReader:
    def __init__(self):
        self.headers = {
            "User-Agent":
            "Mozilla/5.0 (Windows NT 6.3; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/44.0.2403.157 Safari/537.36"
        }

    def convert(self, file_path: str, save_path: str) -> str:
        if file_path.startswith("http"):
            resp = requests.get(file_path, headers=self.headers)
            html_doc = resp.content.decode("utf8")
            soup = BeautifulSoup(html_doc, "html.parser")
            title = soup.find("title").get_text()
            if not title:
                title = soup.find("h1").get_text()
            title = title.strip()
        else:
            html_doc = pnlp.read_file(file_path)
            soup = BeautifulSoup(html_doc, "html.parser")
            title = Path(file_path).stem
        title = title.replace("\n", " ")
        out_fn = title + ".txt"
        out_fp = os.path.join(save_path, out_fn)
        # mydivs = soup.find_all("div", {"class": "Article_content"})
        mydivs = soup.find_all("p")
        has_title = False
        for v in mydivs:
            txt = v.get_text().strip()
            if title_reg.search(txt):
                has_title = True
        res = []
        ch = 1
        for i, v in enumerate(mydivs, start=1):
            txt = v.get_text().strip()
            if not has_title and txt and v.find(
                    "strong") and not title_reg.search(txt):
                zh_ch = pnlp.num_norm.num2zh(ch)
                if ch >= 10:
                    zh_ch = zh_ch[1:]
                txt = f"第{zh_ch}章 {txt}"
                ch += 1
            if txt:
                res.append(txt)
        pnlp.write_file(out_fp, res)
        return out_fn


class DocReader:
    def __init__(self):
        self.doc_name = None

    def convert(
            self,
            doc_file,
            save_path,
            batch=False):

        if isinstance(doc_file, Path):
            doc_file = doc_file.as_posix()
        if batch is False:
            if len(doc_file.strip().split('/')) == 1:
                doc_name = doc_file.strip().split('\\')[-1]
            else:
                doc_name = doc_file.strip().split('/')[-1]
        else:
            doc_name = self.doc_name

        if doc_name.endswith('.docx'):
            out_doc_name = doc_name[:-5]
        elif doc_name.endswith('.doc'):
            out_doc_name = doc_name[:-4]

        if os.path.exists(save_path):
            pass
        else:
            os.mkdir(save_path)

        txt_file = open(
            os.path.join(
                save_path,
                f"{out_doc_name}.txt"),
            'w',
            encoding="utf-8")
        if doc_name.endswith('.doc'):
            doc = subprocess.check_output(['antiword', doc_file]).decode('utf-8')
            txt_file.write(doc)
        elif doc_name.endswith('.docx'):
            file_docx = docx.Document(doc_file)
            for i in range(len(file_docx.paragraphs)):
                txt_file.write("{}\n".format(file_docx.paragraphs[i].text))

        txt_file.close()
        return f"{out_doc_name}.txt"

    def convert_batch(self, doc_dir, save_path):
        file_names = []
        for file_name in os.listdir(doc_dir):
            if len(file_name.split('.')) == 2:
                if (file_name.split('.')[1] == 'doc') or (
                        file_name.split('.')[1] == 'docx'):
                    self.doc_name = file_name.split('.')[0]
                    # print(self.doc_name)
                    file = self.convert(
                        os.path.join(
                            doc_dir,
                            file_name),
                        save_path,
                        batch=True)
                    file_names.append(file)
            else:
                continue
        return file_names


class PDFReader:
    def __init__(self, use_gpu=True):
        """
        两种形式的PDF：
        1、文字版本，直接转成文本格式即可。
        2、图片版本，扫描件。
        图片版本的可能需要OCR识别。较为复杂。直接通过PP-Structure进行识别。
        example:pdf图片格式转成图片。
        pdf_file = 'data/加梯FAQ/2022-02-09沪建设施（2022）97号：加梯管道保护和配合搬迁的通知.pdf'
        doc = fitz.open(pdf_file)
        page_length = doc.page_count
        for i in range(page_length):
            pdf_page = doc.load_page()
            img = pdf_page.get_pixmap()
            img.save("img{}.png".format(i))
        """
        self.recognize = PaddleOCR(use_angle_cls=True,
                                   lang='ch',
                                   use_gpu=use_gpu)

        self.pdf_name = None

    def convert(
            self,
            pdf_path,
            save_path='D:\\AI Projects\\综合文档问答系统\\docfaq\\data\\加梯FAQ',
            batch=False):
        """
        因为不知道输入的pdf里的内容是文本形式的还是扫描件形式的，因此：
        首先要尝试提取文本，如果直接提取文本失败，则PDF就应该是扫描件形式的。
        直接提取失败后尝试使用PPOCR识别提取文本。
        当然，第三种情况：还有可能是文本+扫描件的形式。
        注意输入的save_path的路径。
        """
        if batch is False:
            if len(pdf_path.strip().split('/')) == 1:
                pdf_name = pdf_path.strip().split('\\')[-1][:-4]
            else:
                pdf_name = pdf_path.strip().split('/')[-1][:-4]
        else:
            pdf_name = self.pdf_name

        if os.path.exists(save_path):
            pass
        else:
            os.mkdir(save_path)

        saved_txt = open(
            os.path.join(
                save_path,
                pdf_name +
                '.txt'),
            'w',
            encoding="utf-8")
        # 先确定PDF中内容的形式：
        pdf_doc = pdfplumber.open(pdf_path)
        all_count = 0
        none_page = []  # 直接提取文字得到为空的pdf的页码记录
        all_dict_text = {}  # 建立一个字典key为页码，内容为每一页文本内容。
        for page in pdf_doc.pages:
            all_count += 1
            text = page.extract_text()
            if text == '':
                # 需要确定为空页面还是图片
                none_page.append(all_count)
            else:
                all_dict_text.update({str(all_count): text})
        # 如果直接提取文本为空的计数器和总的页数相等，则说明该PDF全部为扫描件或者为全空白页。
        # 空白页这种情况也要考虑，因此count！=all_count时，
        # 可能有三种情况：1、图片与PDF文本混合的情况 2、全为PDF的情况 3、存在一两页空白页的情况。
        pdf_doc.close()
        doc = fitz.open(pdf_path)
        page_length = doc.page_count
        img_save_path = os.path.join(save_path, 'doc_imgs')
        if os.path.exists(img_save_path):
            pass
        else:
            os.mkdir(img_save_path)

        if len(none_page) == all_count:
            # 说明全部为扫描件。通过PaddleOCR识别。
            print("全为扫描件情况：{}".format(pdf_path))
            texts = ''
            for i in range(page_length):
                pdf_page = doc.load_page(i)
                img = pdf_page.get_pixmap()
                img_path = os.path.join(img_save_path, "img{}.png".format(i))
                img.save(img_path)
                results = self.recognize.ocr(img_path, cls=True)
                for line in results:
                    # print(line)
                    # 对于附件的格式处理。以x坐标排序之后再以y坐标进行排序。
                    texts += line[-1][0]
                    texts += '\n'
                os.remove(img_path)

            saved_txt.write(texts)
            saved_txt.close()
            doc.close()
        else:
            # 说明全部是含文本的内容。直接将all_text写进保存路径下的一个文件即可。
            # 首先处理空白页面内容
            print("混合情况:{}".format(pdf_path))
            for page_lost in none_page:
                print("page_lost:{}".format(page_lost))
                texts = ''
                pdf_page = doc.load_page(page_lost - 1)
                img = pdf_page.get_pixmap()
                img_path = os.path.join(
                    img_save_path, "img{}.png".format(
                        page_lost - 1))
                img.save(img_path)
                results = self.recognize.ocr(img_path, cls=True)
                for line in results:
                    # print(line)
                    texts += line[-1][0]
                    texts += '\n'
                os.remove(img_path)
                all_dict_text.update({str(page_lost): texts})
            # 写所有的txt文件
            for i in range(all_count):
                saved_txt.write(all_dict_text[str(i + 1)])
            saved_txt.close()

        # 删除保存文件的路径。
        os.removedirs(img_save_path)
        return f"{pdf_name}.txt"

    def convert_img(
            self,
            img_path,
            save_path):
        fn = Path(img_path).stem
        texts = ''
        try:
            results = self.recognize.ocr(img_path, cls=True)
            for line in results:
                part = line[-1][0]
                texts += part
                texts += '\n'
        except Exception as e:
            print(f"图片转换失败，错误信息：{e}")
            pass
        # 识别出结果不满10个字，不计为有效识别
        # if len(texts) < 10:
            # return None
        saved_txt = open(
            os.path.join(
                save_path,
                f"{fn}.txt"),
            'w',
            encoding="utf-8")
        saved_txt.write(texts)
        saved_txt.close()
        return f"{fn}.txt"

    def convert_batch(
            self,
            pdf_dir,
            save_path):
        file_names = []
        for file_name in os.listdir(pdf_dir):
            if len(file_name.split('.')) == 2:
                if file_name.split('.')[1] == 'pdf':
                    self.pdf_name = file_name.split('.')[0]
                    file = self.convert(
                        os.path.join(
                            pdf_dir,
                            file_name),
                        save_path,
                        batch=True)
                    file_names.append(file)
            else:
                continue
        return file_names


class PPTReader:

    def __init__(self, use_gpu=True):
        pass

    def convert(self, file_path: str, save_path: str) -> str:
        title = Path(file_path).stem
        out_fn = title + ".txt"
        out_fp = os.path.join(save_path, out_fn)
        with open(out_fp, "w") as f:
            f.write("")
        return out_fn



"""
class PPTReader:
    解决两个问题：
    1、如何读图片。
    2、PP-Structure等OCR识别得到文本。
    def __init__(self, use_gpu=True):
        self.recognize = PaddleOCR(use_angle_cls=True,
                                   lang='ch',
                                   use_gpu=use_gpu)
        self.newpath = None

    Note:PPT文件按页导出成图片，在Windows上可以借助win32com等模块调用PPT软件功能实现。
    但是在linux上不支持pywin32，参考教程：将下述ppt2pdf函数重写即可。
    https://blog.csdn.net/qq_44920726/article/details/104652145
    # todo 同doc的问题一样，主要是linux下的方案

    def ppt2pdf(self, ppt_path):
        self.newpath = os.path.splitext(ppt_path)[0] + '.pdf'
        ppt_app = win32com.client.Dispatch('PowerPoint.Application')
        ppt = ppt_app.Presentations.Open(ppt_path)  # 打开 ppt
        ppt.SaveAs(self.newpath, 32)  # 17数字是转为 ppt 转为图片
        ppt_app.Quit()  # 关闭资源，退出
        os.remove(ppt_path)

    # 关于PPT2TXT的文本格式，以x坐标排序之后再以y坐标进行排序。
    def textformat(self, txt_results):
        pass

    def convert(
            self,
            ppt_path,
            save_path):
        if len(ppt_path.strip().split('/')) == 1:
            ppt_name = ppt_path.strip().split('\\')[-1]
        else:
            ppt_name = ppt_path.strip().split('/')[-1]

        if ppt_name.endswith('.pptx'):
            ppt_name = ppt_name[:-5]
        elif ppt_name.endswith('.ppt'):
            ppt_name = ppt_name[:-4]

        self.ppt2pdf(ppt_path)
        doc = fitz.open(self.newpath)

        page_length = doc.page_count
        img_save_path = os.path.join(save_path, 'ppt_imgs')
        if os.path.exists(img_save_path):
            pass
        else:
            os.mkdir(img_save_path)

        saved_txt = open(
            os.path.splitext(ppt_path)[0] +
            '.txt',
            "w",
            encoding="utf-8")
        texts = ''
        for i in range(page_length):
            pdf_page = doc.load_page(i)
            img = pdf_page.get_pixmap()
            img_path = os.path.join(img_save_path, "img{}.png".format(i))
            img.save(img_path)
            results = self.recognize.ocr(img_path, cls=True)
            for line in results:
                # print(line)
                texts += line[-1][0]
                texts += '\t'
            texts += '\n'
            os.remove(img_path)
        saved_txt.write(texts)
        saved_txt.close()
        os.removedirs(img_save_path)
        return f"{ppt_name}.txt"
"""

if __name__ == "__main__":
    # 注意，输入的文件路径
    """
    convert_batch:支持直接将文件夹下的所有PDF文件或者DOC、DOCX文件转换成文本文件。
    convert：需传入要转换文件的路径。
    convert图片的时候，使用PDFReader里的convert_img，PDFMode要设置为False并给出文档命名，同时输入保存地址。
    """
    pdf_file = 'D:\\AI Projects\\综合文档问答系统\\docfaq\\data\\加梯FAQ\\沪公积金〔2021〕59号.pdf'
    doc_file = 'D:\\AI Projects\\综合文档问答系统\\docfaq\\data\\加梯FAQ\\静安区既有多层住宅加装电梯项目安全质量管理(报审稿修)12.2.doc'
    img_path = "D:\\AI Projects\\综合文档问答系统\\docfaq\\data\\加梯FAQ\\doc_imgs\\img0.png"
    pdf_batch = 'D:\\AI Projects\\综合文档问答系统\\docfaq\\data\\加梯FAQ'
    # ppt_path = 'D:\\AI Projects\\综合文档问答系统\\docfaq\\data\\加梯FAQ\\第0章绪论.ppt'

    # doc_reader = DocReader()
    # doc_reader.convert(doc_file, "D:\\AI Projects\\综合文档问答系统\\docfaq\\data\\加梯FAQ")
    # doc_reader.convert_batch(doc_file)

    # pdf_reader = PDFReader()
    # print("###############PDF file converting#############")
    # pdf_reader.convert(pdf_file, "D:\\AI Projects\\综合文档问答系统\\docfaq\\data\\加梯FAQ")
    # print("###############Img convrting################")
    # pdf_reader.convert_img(img_path)
    # print("##############Batch converting##############")
    # pdf_reader.convert_batch(pdf_batch)

    # ppt_reader = PPTReader()
    # ppt_reader.convert(ppt_path)
